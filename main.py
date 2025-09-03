from docx import Document
from tqdm import tqdm
import tkinter as tk
from tkinter import filedialog
import time
from docx2pdf import convert
import os

def get_files():
    # Esconde a janela principal do tkinter
    root = tk.Tk()
    root.withdraw()  # Esconde a janela principal
    
    # Abre a janela para seleção de arquivo
    arquivos = filedialog.askopenfilenames(
        title="Selecione um arquivo",
        filetypes=[
            ("Documento Word", "*.docx"),
        ]
    )

    return arquivos

def get_fields(arquivos):
    campos = []
    for arquivo in arquivos:
        document = Document(arquivo)
        for paragraph in document.paragraphs:
            # Pega o texto completo do parágrafo
            texto_completo = paragraph.text.strip()
            
            # Procura por padrões *TEXTO* no texto completo
            import re
            matches = re.findall(r'\*[^*]+\*', texto_completo)
            for match in matches:
                if match not in campos:
                    campos.append(match)
    return campos

def write_file(campos_preenchidos, nome_arquivos):
    for arquivo, nome_arquivo in nome_arquivos.items():
        document = Document(arquivo)

        for key, value in tqdm(campos_preenchidos.items()):
            for paragraph in document.paragraphs:
                # Junta o texto de todos os runs
                texto_runs = ''.join(run.text for run in paragraph.runs)
                if key in texto_runs:
                    # Substitui o campo no texto total
                    novo_texto = texto_runs.replace(key, value)
                    # Distribui o novo texto nos runs, preservando a formatação
                    i = 0
                    for run in paragraph.runs:
                        run_len = len(run.text)
                        run.text = novo_texto[i:i+run_len]
                        i += run_len
        document.save(nome_arquivo)
        print(f"{nome_arquivo} finalizado!")

def converter_com_docx2pdf(docx_path, pdf_path):
    try:
        convert(docx_path, pdf_path)
        return True
    except Exception as e:
        print(f"Erro: {e}")
        return False


def main():
    print('''+--------------------------+
|  Preenchendo Documentos  |
+--------------------------+\n''')
    
    # Escolhendo arquivos
    print("1. Selecione os arquivos...")
    arquivos = get_files()
    for arquivo in arquivos:
        print(f"- {arquivo.split('/')[len(arquivo.split('/')) - 1]}")

    time.sleep(1)
    print()
    
    # Procurando quais campos serão preenchidos
    print("2. Lendo os arquivos...")
    campos = get_fields(arquivos)
    print("Os campos que serão preenchidos são:")
    for campo in campos:
        print(f"- {campo[1:len(campo) - 1]}")

    time.sleep(1)
    print()

    # Recendo input dos dados
    print("3. Insira os dados a serem preenchidos...")
    campos_preenchidos = {}
    for campo in campos:
        input_user = input(f"{campo[1:len(campo) - 1]}: ")
        campos_preenchidos[campo] = input_user

    time.sleep(1)
    print()

    # Nome dos arquivos
    print("4. Quais serão os nomes dos arquivos?")
    nome_arquivos = {}
    for arquivo in arquivos:
        nome_atual = os.path.basename(arquivo)
        caminho = os.path.dirname(arquivo)
        novo_nome = input(f"- {nome_atual}: ")
        nome_final = os.path.join(caminho, f"{novo_nome}.docx")
        nome_arquivos[arquivo] = nome_final

    time.sleep(1)
    print()

    # Escrevendo no documento os dados
    print("5. Escrevendo dados no documento...")
    write_file(campos_preenchidos, nome_arquivos)
    
    time.sleep(1)
    print()

    # Converter para pdf
    converter_pdf = True if input("6. Deseja converter para PDF? (Responda com sim ou não)\n>> ").lower() == 'sim' else False
    if converter_pdf:
        for arquivo, nome_arquivo in nome_arquivos.items():
            pdf_path = os.path.splitext(nome_arquivo)[0] + ".pdf"
            sucesso = converter_com_docx2pdf(nome_arquivo, pdf_path)
            if sucesso:
                print(f"PDF gerado: {pdf_path}")
            else:
                print(f"Falha ao converter {nome_arquivo} para PDF.")
    
    print()
    time.sleep(3)

    print("Os arquivos foram salvos com sucesso!\nEncerrando o programa.")
    time.sleep(5)

    
if __name__ == "__main__":
    main()
